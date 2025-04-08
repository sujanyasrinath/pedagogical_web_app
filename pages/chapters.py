from flask import Blueprint, flash, render_template, request, redirect, send_file, session
from docx import Document
from docx.shared import Pt

from docx.shared import Inches
from .utils.dataset import Dataset

import matplotlib
matplotlib.use("Agg")

import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
import os
import io
import base64
from io import StringIO
import random
from rapidfuzz import process, fuzz
import re

import matplotlib.pyplot as plt
import seaborn as sns
import os
from io import BytesIO
from docx.shared import Inches
from scipy.stats import trim_mean, mode



def save_plot_image(plot_func, filename, df_column, **kwargs):
    """Helper function to generate a plot and save it as an image file."""
    fig, ax = plt.subplots(figsize=(6, 4))
    plot_func(data=df_column, ax=ax, **kwargs)  # ✅ Pass ax as keyword argument
    
    img_path = os.path.join("static", "plots", filename)
    os.makedirs(os.path.dirname(img_path), exist_ok=True)  # Ensure directory exists
    plt.savefig(img_path, format="png", bbox_inches="tight")
    plt.close(fig)

    return img_path

chapters = Blueprint("chapters/", __name__, url_prefix="/chapters")

ALL_CHAPTERS = [
    {
        "index": i,
        "description": f"Description for Chapter {i}",
        "url": f'/chapters/{i}'
    } for i in range(1, 11)
]

def create_chapter_1_document(doc, selected_dataset, chapter_1_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    course = chapter_1_details.get('course')
    question = chapter_1_details.get('question')

    if show_answer:
        doc.add_heading(f'{course} Chapter 1 Answer Key', level=1)
    else:
        doc.add_heading(f'{course} Chapter 1 Assignment', level=1) #add in course name from input
        doc.add_paragraph('Student Name:')
        doc.add_paragraph('Date:')

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

    # doc.add_heading('Section 1: Research Question', level=2)
    # doc.add_paragraph("Answer the following research question")
    # doc.add_paragraph(question)

    # doc.add_heading('Section 2: Metadata File', level=2)
    # doc.add_paragraph("Review the partially completed metadata table below. Fill in the missing pieces based on your understanding of the dataset.")
    # #add manipulated metadata file from input dataset here

    # doc.add_heading('Section 3: STAR Framework - Dataset Overview', level=2)
    # doc.add_paragraph("Use the STAR framework to describe the dataset and its context. Write a short paragraph for each component.")
    # #add dataset provided from input here

    # Section 1
    doc.add_heading('Section 1: Research Question', level=2)

    instruction_para = doc.add_paragraph()
    instruction_run = instruction_para.add_run("Answer the following research question")
    instruction_run.italic = True

    response_para = doc.add_paragraph()
    response_run = response_para.add_run(question)
    response_run.bold = False
    response_run.font.size = Pt(12)

    if show_answer:
        doc.add_paragraph("Custom grading per student")

    # Section 2
    doc.add_heading('Section 2: Metadata File', level=2)
    metadata_para = doc.add_paragraph()
    metadata_run = metadata_para.add_run("Review the partially completed metadata table below. Fill in the missing pieces based on your understanding of the dataset.")
    metadata_run.italic = True

    # inserting a table here for metadata
#     metadata_table = doc.add_table(rows=1, cols=5)
#     metadata_table.style = 'Table Grid'
#     hdr_cells = metadata_table.rows[0].cells
#     headers = ['Variable Name', 'Description', 'Data Type', 'Example Value', 'Units / Notes']
#     for i, header in enumerate(headers):
#         hdr_cells[i].text = header

# # Add rows from dataset with some missing values for students to complete
#     for col in df.columns:
#         row_cells = metadata_table.add_row().cells
#         row_cells[0].text = col  # Variable Name
#         row_cells[1].text = ''   # Description - leave blank
#         row_cells[2].text = str(df[col].dtype)  # Data Type
#         example_value = df[col].dropna().iloc[0] if not df[col].dropna().empty else ''
#         row_cells[3].text = str(example_value)  # Example Value
#         row_cells[4].text = ''   # Units / Notes - leave blank
    metadata_table = doc.add_table(rows=1, cols=5)
    metadata_table.style = 'Table Grid'
    hdr_cells = metadata_table.rows[0].cells
    headers = ['Variable Name', 'Description', 'Data Type', 'Example Value', 'Units / Notes']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

# Prepare initial data rows (fill in everything first)
    rows_data = []
    for col in df.columns:
        example_value = df[col].dropna().iloc[0] if not df[col].dropna().empty else ''
        rows_data.append([
            col,                         # Variable Name (always filled)
            'Description goes here',     # Description (can be blanked)
            str(df[col].dtype),          # Data Type (can be blanked)
            str(example_value),          # Example Value (can be blanked)
            'Units or notes here'        # Units / Notes (can be blanked)
        ])

    # Create list of cell indices to potentially blank (row index, column index)
    # Only blank columns 1–4 (i.e., skip Variable Name)
    editable_indices = [(i, j) for i in range(len(rows_data)) for j in range(1, 5)]

    # Blank about 30% of all editable cells
    num_to_blank = int(len(editable_indices) * 0.3)
    blank_indices = random.sample(editable_indices, num_to_blank)

    # Apply the blanking
    if not show_answer:
        for i, j in blank_indices:
            rows_data[i][j] = ''

    # Add rows to the Word table
    for row in rows_data:
        row_cells = metadata_table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = val

    # Section 3
    doc.add_heading('Section 3: STAR Framework', level=2)
    star_para = doc.add_paragraph()
    star_run = star_para.add_run("Describe the results of the tasks using the STAR Framework")
    star_run.italic = True

    if show_answer:
        doc.add_paragraph("Custom grading per student")

    doc.add_heading('Download the Dataset', level=2)
    doc.add_paragraph(
        "You can download the dataset using the following link:"
    )

    dataset_url = f"{request.host_url}static/uploads/{selected_dataset}"

    # Replace this with your actual download link logic
    doc.add_paragraph(dataset_url)


def create_chapter_2_document(doc, selected_dataset, chapter_2_details, show_answer=False):
    print("creating chapter 2 doc")
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    course = chapter_2_details.get('course')
    task_3 = chapter_2_details.get('task3')
    task_4 = chapter_2_details.get('task4')
    task_5 = chapter_2_details.get('task5')

    def normalize(text):
        return re.sub(r'[^\w\s]', '', text).strip().lower()  # remove punctuation, trim, lower
        



    if show_answer:
        doc.add_heading(f'{course} Chapter 2 Answer Key', level=1)
        doc.add_heading("Final Cleaned Dataset with All Tasks Applied", level=2)
        final_df = df.copy()
    else:
        doc.add_heading(f'{course} Chapter 2 Assignment', level=1) #add in course name from input
        doc.add_paragraph('Student Name:')
        doc.add_paragraph('Date:')

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')
    if not show_answer:
        para = doc.add_paragraph()
        run = para.add_run("Please submit a csv file with the manipulated dataset along with the assignment document.")
        run.italic = True

    # Section 1
    if not show_answer:

        doc.add_heading('Section 1: Removing duplicate rows', level=2)
        doc.add_paragraph('The following dataset has some duplicate rows. Remove these duplicate rows and give the number of rows in the original dataset.')
        
        # Duplicate 25 random rows
        duplicated_df = pd.concat([df, df.sample(25, replace=True)])
        duplicated_path = os.path.join("static", "csv", f"duplicated_{selected_dataset}")
        original_path = os.path.join("static", "csv", f"original_{selected_dataset}")

        print("line 199")
        duplicated_df.to_csv(duplicated_path, index=False)
        df.to_csv(original_path, index=False)
        print("line 203")


        doc.add_paragraph("Download the dataset with duplicated rows:")

        # ✅ Plain text file paths instead of hyperlinks
        doc.add_paragraph(f"📂 Duplicated Dataset: {request.host_url}{duplicated_path}")
    
    else:
        # Task 1: Remove duplicates
        final_df = final_df.drop_duplicates()
        original_count = len(df)
        doc.add_paragraph(f"Section 1\nNumber of rows in original dataset: {original_count}")

    # Section 2
    if not show_answer:
        doc.add_heading('Section 2: Replacing missing values', level=2)
        doc.add_paragraph('Please find the dataset with missing values below. Fill in the missing values of the numeric columns using the mean imputation method.')

        # Get only numeric columns
        numeric_cols = df.select_dtypes(include=['number']).columns

        # Create a copy of the DataFrame for missing value injection
        df_with_missing = df.copy()

        # Set 10% of each numeric column's values to NaN
        for col in numeric_cols:
            n = len(df_with_missing)
            n_missing = int(n * 0.1)
            missing_indices = df_with_missing.sample(n=n_missing).index
            df_with_missing.loc[missing_indices, col] = None

        # Save the dataset with missing values
        missing_path = os.path.join("static", "csv", f"missing_{selected_dataset}")
        df_with_missing.to_csv(missing_path, index=False)

        # Add download instruction to doc
        doc.add_paragraph("Download the dataset with missing values:")

        # Generate full URL to make it downloadable (just like duplicated_path)
        download_url_missing = os.path.join(request.host_url, "static", "csv", f"missing_{selected_dataset}")
        doc.add_paragraph(f"📂 Missing Values Dataset: {download_url_missing}")
    else:
        for col in final_df.select_dtypes(include='number').columns:
            n_missing = int(len(final_df) * 0.1)
            missing_indices = final_df.sample(n=n_missing).index
            final_df.loc[missing_indices, col] = None
            mean_val = final_df[col].mean()
            final_df[col].fillna(mean_val, inplace=True)
        doc.add_paragraph("Section 2\nThe dataset shown above has the imputed values in italics. Grade based on comparison between student dataset and dataset shown above.")


    # Section 3
    if not show_answer:
        doc.add_heading('Section 3: Clean Incorrect Formats', level=2)
        doc.add_paragraph('Below are the replacement rules that are provided. Apply these cleaning rules to the dataset to standardize the values in the selected column. Once this task is done answer the following questions')

        # fuzzy matching to get pairs
        # Get the column the professor wants to work with
        # column_to_clean = df[task_3].dropna().astype(str)

        # # Get unique values in the column
        # unique_values = column_to_clean.unique()

        # # Set a similarity threshold
        # SIMILARITY_THRESHOLD = 85

        # # Generate groups of similar values
        # corrections = {}
        # for value in unique_values:
        #     # Skip if this value is already a "correction"
        #     if value in corrections.values():
        #         continue
        #     matches = process.extract(value, unique_values, scorer=fuzz.ratio, limit=10)
        #     for match, score, _ in matches:
        #         if match != value and score >= SIMILARITY_THRESHOLD:
        #             corrections[match] = value  # match → correct value

        # # Create a DataFrame for the correction pairs
        # corrections_df = pd.DataFrame(list(corrections.items()), columns=["Incorrect Text", "Corrected Text"])

        # # Add the table to the Word doc
        # if not corrections_df.empty:
        #     table = doc.add_table(rows=1, cols=2)
        #     table.style = 'Table Grid'
        #     hdr_cells = table.rows[0].cells
        #     hdr_cells[0].text = 'Incorrect Text'
        #     hdr_cells[1].text = 'Corrected Text'

        #     for _, row in corrections_df.iterrows():
        #         row_cells = table.add_row().cells
        #         row_cells[0].text = str(row["Incorrect Text"])
        #         row_cells[1].text = str(row["Corrected Text"])
        # else:
        #     doc.add_paragraph("No similar entries were found to suggest corrections.")
        column_to_clean = task_3  # e.g., "country"
        df["normalized_column"] = df[column_to_clean].dropna().astype(str).apply(normalize)

        norm_to_originals = {}
        for val in df[column_to_clean].dropna().astype(str):
            norm = normalize(val)
            if norm not in norm_to_originals:
                norm_to_originals[norm] = set()
            norm_to_originals[norm].add(val)

        # value_counts = df[task_3].dropna().astype(str).value_counts()
        value_counts = df["normalized_column"].value_counts()
        MIN_FREQ_FOR_CORRECT = 2 #max(3, int(len(df) * 0.01))  # ≥3 or ≥1% of dataset size
        MAX_FREQ_FOR_INCORRECT = 2  # typo-like entries usually occur only once
        FUZZY_THRESHOLD = 75
        correct_values = value_counts[value_counts >= MIN_FREQ_FOR_CORRECT].index.tolist()
        suspect_values = value_counts[value_counts <= MAX_FREQ_FOR_INCORRECT].index.tolist()

        # corrections = {}

        # # for suspect in suspect_values:
        # #     match, score, _ = process.extractOne(suspect, correct_values, scorer=fuzz.ratio)
        # #     if score >= FUZZY_THRESHOLD:
        # #         corrections[suspect] = match

        # for suspect in suspect_values:
        #     result = process.extractOne(suspect, correct_values, scorer=fuzz.ratio)
        #     if result:
        #         match, score, _ = result
        #         if score >= FUZZY_THRESHOLD:
        #             corrections[suspect] = match
        
        # corrections_df = pd.DataFrame(list(corrections.items()), columns=["Incorrect Text", "Corrected Text"])
        #     # Check if corrections_df has any rows
        # if not corrections_df.empty:
        #     # Create a table in the Word doc with header row
        #     table = doc.add_table(rows=1, cols=len(corrections_df.columns))
        #     table.style = 'Table Grid'

        #     # Set headers
        #     hdr_cells = table.rows[0].cells
        #     for idx, col_name in enumerate(corrections_df.columns):
        #         hdr_cells[idx].text = col_name

        #     # Add data rows
        #     for _, row in corrections_df.iterrows():
        #         row_cells = table.add_row().cells
        #         for idx, value in enumerate(row):
        #             row_cells[idx].text = str(value)
        # else:
        #     doc.add_paragraph("No replacement pairs were detected.")
        corrections = {}
        for suspect in suspect_values:
            result = process.extractOne(suspect, correct_values, scorer=fuzz.ratio)
            if result:
                match, score, _ = result
                if score >= FUZZY_THRESHOLD:
                    corrections[suspect] = match

        # Step 6: Rebuild correction pairs using original messy values
        replacement_rows = []
        for norm_incorrect, norm_correct in corrections.items():
            original_variants = norm_to_originals.get(norm_incorrect, {norm_incorrect})
            corrected_variant = next(iter(norm_to_originals.get(norm_correct, {norm_correct})))
            for original in original_variants:
                replacement_rows.append((original, corrected_variant))

        corrections_df = pd.DataFrame(replacement_rows, columns=["Incorrect Text", "Corrected Text"])

        # Step 7: Add to Word doc
        if not corrections_df.empty:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Add headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Incorrect Text"
            hdr_cells[1].text = "Corrected Text"

            # Add data rows
            for _, row in corrections_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row["Incorrect Text"])
                row_cells[1].text = str(row["Corrected Text"])
        else:
            doc.add_paragraph("No replacement pairs were detected.")

    else:
        text_col = task_3
        final_df[text_col] = final_df[text_col].astype(str).apply(normalize)
        doc.add_paragraph("Section 3\nThe dataset shown above has the incorrect formats cleaned. Grade based on comparison between student dataset and dataset shown above.")


    # Section 4
    if not show_answer:
        doc.add_heading(' Section 4: Creating a new categorical column', level=2)
        para = doc.add_paragraph()
        para.add_run("Create a new categorical variable that contains a categorized version of the numeric variable ").italic = True
        para.add_run(f"'{task_4}'").bold = True
        para.add_run(". Please specify the thresholds and what each range should be categorized as. The first few rows are shown here").italic = True

        unique_values = df[task_4].dropna().astype(str).unique()[:10]

        # Add heading and description
        doc.add_paragraph(f"Values from column: '{task_4}'", style='Heading 3')

        # Create a table with one column
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = task_4  # header

        # Add values to the table
        for val in unique_values:
            row_cells = table.add_row().cells
            row_cells[0].text = str(val)
    else:
        cat_col = task_4
        thresholds = pd.qcut(final_df[cat_col], 3, labels=["Low", "Medium", "High"])
        final_df[f"{cat_col}_category"] = thresholds
        doc.add_paragraph("Section 4\nThe dataset shown above has the categorized values in italics. Grade based on comparison between student dataset and dataset shown above.")

    # Section 5
    if not show_answer:
        doc.add_heading('Section 5: Create a Boolean Variable that Represents a Categorical Variable', level=2)
        para = doc.add_paragraph()
        para.add_run("Please review the column").italic = True
        para.add_run(f"'{task_5}'").bold = True
        para.add_run("below. Please find a few rows here").italic = True
        unique_values = df[task_5].dropna().astype(str).unique()[:10]

        # Add heading and description
        doc.add_paragraph(f"Values from column: '{task_5}'", style='Heading 3')

        # Create a table with one column
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = task_5  # header

        # Add values to the table
        for val in unique_values:
            row_cells = table.add_row().cells
            row_cells[0].text = str(val)

        column_name = task_5  
        most_common_value = df[column_name].dropna().astype(str).value_counts().idxmax()

        rule_1 = f"True if {column_name} == '{most_common_value}'"
        rule_2 = f"False if {column_name} != '{most_common_value}'"

        # Add to Word doc
        doc.add_heading("Create a Boolean Variable", level=2)

        doc.add_paragraph(
            f"Create a Boolean variable where:\n"
            f"- Rule 1: {rule_1}\n"
            f"- Rule 2: {rule_2}\n"
        )

        
        doc.add_paragraph(f"Based on the new column answer the following questions")
        doc.add_paragraph(f"How many of 'True' are in the new column?")
        doc.add_paragraph(f"How many of 'False' are in the new column?")
    else:
        bool_col = task_5
        common_val = final_df[bool_col].value_counts().idxmax()
        final_df[f"{bool_col}_is_{common_val}"] = final_df[bool_col] == common_val
        x_count = (final_df[bool_col] == common_val).sum()
        y_count = len(final_df) - x_count

        doc.add_paragraph("Section 5\nThe dataset shown above has a new column representing a boolean variable. Grade based on comparison between student dataset and dataset shown above.")
        doc.add_paragraph(f"Number of times '{common_val}' appeared before cleaning: {x_count}")
        doc.add_paragraph(f"Number of times other values appeared: {y_count}")
    if show_answer:
        # doc.add_heading("Final Cleaned Dataset with All Tasks Applied", level=2)
        # final_df = df.copy()

        # # Task 1: Remove duplicates
        # final_df = final_df.drop_duplicates()
        # original_count = len(df)
        # doc.add_paragraph(f"Section 1\nNumber of rows in original dataset: {original_count}")

        # # Task 2: Impute mean into 10% missing numeric values
        # for col in final_df.select_dtypes(include='number').columns:
        #     n_missing = int(len(final_df) * 0.1)
        #     missing_indices = final_df.sample(n=n_missing).index
        #     final_df.loc[missing_indices, col] = None
        #     mean_val = final_df[col].mean()
        #     final_df[col].fillna(mean_val, inplace=True)
        # doc.add_paragraph("Section 2\nThe dataset shown above has the imputed values in italics. Grade based on comparison between student dataset and dataset shown above.")

        # # Task 3: Clean incorrect text formats
        # text_col = task_3
        # final_df[text_col] = final_df[text_col].astype(str).apply(normalize)
        # doc.add_paragraph("Section 3\nThe dataset shown above has the incorrect formats cleaned. Grade based on comparison between student dataset and dataset shown above.")

        # # Task 4: Categorize numeric column using auto-threshold
        # cat_col = task_4
        # thresholds = pd.qcut(final_df[cat_col], 3, labels=["Low", "Medium", "High"])
        # final_df[f"{cat_col}_category"] = thresholds
        # doc.add_paragraph("Section 4\nThe dataset shown above has the categorized values in italics. Grade based on comparison between student dataset and dataset shown above.")

        # # Task 5: Boolean from most common value
        # bool_col = task_5
        # common_val = final_df[bool_col].value_counts().idxmax()
        # final_df[f"{bool_col}_is_{common_val}"] = final_df[bool_col] == common_val
        # x_count = (final_df[bool_col] == common_val).sum()
        # y_count = len(final_df) - x_count

        # doc.add_paragraph("Section 5\nThe dataset shown above has a new column representing a boolean variable. Grade based on comparison between student dataset and dataset shown above.")
        # doc.add_paragraph(f"Number of times '{common_val}' appeared before cleaning: {x_count}")
        # doc.add_paragraph(f"Number of times other values appeared: {y_count}")

        # Save cleaned dataset
        final_path = os.path.join("static", "csv", f"chapter2_answer_{selected_dataset}")
        final_df.to_csv(final_path, index=False)

        doc.add_paragraph("Download the fully processed dataset used to generate the answer key:")
        doc.add_paragraph(f"{request.host_url}{final_path}")


def create_chapter_3_document(doc, selected_dataset, chapter_3_details, show_answer=False):
    """Generates a Word document for Chapter 3 with visualizations."""

    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)


    course = chapter_3_details.get('course')
    ch3_x = chapter_3_details.get('selectingx')
    ch3_y = chapter_3_details.get('selectingy')
    task_3 = chapter_3_details.get('task3')
    task_4 = chapter_3_details.get('task4')
    task_5 = chapter_3_details.get('task5')

    print(chapter_3_details)
    def normalize(text):
        return re.sub(r'[^\w\s]', '', text).strip().lower()  # remove punctuation, trim, lower
    

    doc.add_heading("Final Cleaned Dataset with All Tasks Applied", level=2)
    final_df = df.copy()

    # Task 1: Remove duplicates
    final_df = final_df.drop_duplicates()

    # Task 2: Impute mean into 10% missing numeric values
    for col in final_df.select_dtypes(include='number').columns:
        n_missing = int(len(final_df) * 0.1)
        missing_indices = final_df.sample(n=n_missing).index
        final_df.loc[missing_indices, col] = None
        mean_val = final_df[col].mean()
        final_df[col].fillna(mean_val, inplace=True)
    
    # Task 3: Clean incorrect text formats
    text_col = task_3
    final_df[text_col] = final_df[text_col].astype(str).apply(normalize)
    doc.add_paragraph("Section 3\nThe dataset shown above has the incorrect formats cleaned. Grade based on comparison between student dataset and dataset shown above.")

    # Task 4: Categorize numeric column using auto-threshold
    cat_col = task_4
    thresholds = pd.qcut(final_df[cat_col], 3, labels=["Low", "Medium", "High"])
    final_df[f"{cat_col}_category"] = thresholds
    doc.add_paragraph("Section 4\nThe dataset shown above has the categorized values in italics. Grade based on comparison between student dataset and dataset shown above.")

    # Task 5: Boolean from most common value
    bool_col = task_5
    common_val = final_df[bool_col].value_counts().idxmax()
    final_df[f"{bool_col}_is_{common_val}"] = final_df[bool_col] == common_val
    x_count = (final_df[bool_col] == common_val).sum()
    y_count = len(final_df) - x_count

    doc.add_paragraph("Section 5\nThe dataset shown above has a new column representing a boolean variable. Grade based on comparison between student dataset and dataset shown above.")
    doc.add_paragraph(f"Number of times '{common_val}' appeared before cleaning: {x_count}")
    doc.add_paragraph(f"Number of times other values appeared: {y_count}")

    # Save cleaned dataset
    final_path = os.path.join("static", "csv", f"chapter2_answer_{selected_dataset}")
    final_df.to_csv(final_path, index=False)       

    if show_answer:
            doc.add_heading(f'{course} Chapter 3 Answer Key', level=1)
    else:
        doc.add_heading(f'{course} Chapter 3 Assignment', level=1) #add in course name from input
        doc.add_paragraph('Student Name:')
        doc.add_paragraph('Date:')

        doc.add_paragraph(f'Selected Dataset: {final_df}')
        if not show_answer:
            para = doc.add_paragraph()
            run = para.add_run("Please submit a csv file with the manipulated dataset along with the assignment document.")
            run.italic = True
            
    # plot tasks
    if not show_answer:
        doc.add_paragraph('here are the plots and here are the variables to use for each plot. please submit a file with these plots added')
        doc.add_paragraph(
        'Here are the plots and the variables to use for each plot. Please submit a file with these plots added.')

        # Extract plots and variable names
        plot_types = chapter_3_details.getlist("plot_types") if hasattr(chapter_3_details, "getlist") else chapter_3_details.get("plot_types", [])
        if isinstance(plot_types, str):
            plot_types = [plot_types]

        x = ch3_x
        y = ch3_y

        # Add ordered list
        doc.add_paragraph("Selected Plots and Variables:", style='Heading 3')

        for plot_type in plot_types:
            if "with a Categorical Variable" in plot_type or "Scatterplot" in plot_type or "Bubble Chart" in plot_type:
                desc = f"{plot_type} — X: {x}, Y: {y}"
            elif plot_type in ["Box Plot", "Violin Plot"]:
                desc = f"{plot_type} — Y: {y}"
            elif plot_type in ["Histogram", "Density Plot", "Bar Chart"]:
                desc = f"{plot_type} — X: {x}"
            else:
                desc = f"{plot_type} — X: {x}, Y: {y}"  # fallback

            doc.add_paragraph(desc, style='List Number')
    else:
        # details of plots
        plot_types = chapter_3_details.getlist("plot_types") if hasattr(chapter_3_details, "getlist") else chapter_3_details.get("plot_types", [])
        if isinstance(plot_types, str):
            plot_types = [plot_types]

        print(plot_types, chapter_3_details.get('plot_types'))

        # Retrieve selected x and y variables (ch3_x, ch3_y)
        x = ch3_x
        y = ch3_y

        # Create a directory to store plots if not exists
        plot_dir = "static/plots"
        os.makedirs(plot_dir, exist_ok=True)

        doc.add_heading("Section 6: Plots", level=2)

        # Loop over each selected plot type and generate a plot
        for plot_type in plot_types:
            fig, ax = plt.subplots()

            if plot_type == "Histogram":
                sns.histplot(data=final_df, x=x, ax=ax)
            elif plot_type == "Density Plot":
                sns.kdeplot(data=final_df, x=x, ax=ax)
            elif plot_type == "Box Plot":
                sns.boxplot(data=final_df, y=y, ax=ax)
            elif plot_type == "Violin Plot":
                sns.violinplot(data=final_df, y=y, ax=ax)
            elif plot_type == "Bar Chart":
                sns.countplot(data=final_df, x=x, ax=ax)
            elif plot_type == "Box Plot with a Categorical Variable":
                sns.boxplot(data=final_df, x=x, y=y, ax=ax)
            elif plot_type == "Violin Plot with a Categorical Variable":
                sns.violinplot(data=final_df, x=x, y=y, ax=ax)
            elif plot_type == "Scatterplot":
                sns.scatterplot(data=final_df, x=x, y=y, ax=ax)
            elif plot_type == "Bubble Chart":
                sns.scatterplot(data=final_df, x=x, y=y, size=y, sizes=(20, 200), alpha=0.5, ax=ax)
            else:
                continue  # Skip unknown plot types

            ax.set_title(plot_type)
            plt.tight_layout()
            # Generate a file name (e.g., histogram_dataset.png)
            plot_filename = f"{plot_type.replace(' ', '_').lower()}_{selected_dataset}.png"
            full_path = os.path.join(plot_dir, plot_filename)
            fig.savefig(full_path)
            plt.close(fig)

            # Add a paragraph in the doc to label this plot
            doc.add_paragraph(plot_type)
            # Add the plot image to the Word doc
            doc.add_picture(full_path, width=Inches(5.5))
        # --- End Plot Section ---   




def create_chapter_4_document(doc, selected_dataset, chapter_4_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)


    course = chapter_4_details.get('course')
    task_3 = chapter_4_details.get('task3')
    task_4 = chapter_4_details.get('task4')
    task_5 = chapter_4_details.get('task5')
    ch4_t1 = chapter_4_details.get('ch4_t1')
    ch4_scatter_x = chapter_4_details.get('ch4_scatter_x')
    ch4_scatter_y = chapter_4_details.get('ch4_scatter_y')
    ch4_column_x = chapter_4_details.get('ch4_column_x')
    ch4_column_y = chapter_4_details.get('ch4_column_y')


    def normalize(text):
        return re.sub(r'[^\w\s]', '', text).strip().lower()  # remove punctuation, trim, lower
    

    #doc.add_heading("Final Cleaned Dataset with All Tasks Applied", level=2)
    final_df = df.copy()

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

    # Task 1: Remove duplicates
    final_df = final_df.drop_duplicates()

    # Task 2: Impute mean into 10% missing numeric values
    for col in final_df.select_dtypes(include='number').columns:
        n_missing = int(len(final_df) * 0.1)
        missing_indices = final_df.sample(n=n_missing).index
        final_df.loc[missing_indices, col] = None
        mean_val = final_df[col].mean()
        final_df[col].fillna(mean_val, inplace=True)
    
    # Task 3: Clean incorrect text formats
    text_col = task_3
    final_df[text_col] = final_df[text_col].astype(str).apply(normalize)
    #doc.add_paragraph("Section 3\nThe dataset shown above has the incorrect formats cleaned. Grade based on comparison between student dataset and dataset shown above.")

    # Task 4: Categorize numeric column using auto-threshold
    cat_col = task_4
    thresholds = pd.qcut(final_df[cat_col], 3, labels=["Low", "Medium", "High"])
    final_df[f"{cat_col}_category"] = thresholds
    #doc.add_paragraph("Section 4\nThe dataset shown above has the categorized values in italics. Grade based on comparison between student dataset and dataset shown above.")

    # Task 5: Boolean from most common value
    bool_col = task_5
    common_val = final_df[bool_col].value_counts().idxmax()
    final_df[f"{bool_col}_is_{common_val}"] = final_df[bool_col] == common_val
    
    # Save cleaned dataset
    final_path = os.path.join("static", "csv", f"chapter2_answer_{selected_dataset}")
    final_df.to_csv(final_path, index=False)       

    if show_answer:
            doc.add_heading(f'{course} Chapter 4 Answer Key', level=1)
    else:
        doc.add_heading(f'{course} Chapter 4 Assignment', level=1) #add in course name from input
        doc.add_paragraph('Student Name:')
        doc.add_paragraph('Date:')

        doc.add_paragraph(f'Selected Dataset: {final_df}')

    # --- CHAPTER 4 TASKS ---

    # Task 1: Summary stats after outlier imputation
    doc.add_heading('Task 1: Summary Statistics After Outlier Handling', level=2)
    if not show_answer:
        doc.add_paragraph(f"Variable selected: {ch4_t1}")
        values = df[ch4_t1].dropna().astype(float).head(10)
        for val in values:
            doc.add_paragraph(str(val))
        doc.add_paragraph("Please calculate and interpret the following after handling outliers: mean, mode, median, 90% trimmed mean, range, variance, and standard deviation.")
    else:
        # Inject outliers
        final_df[ch4_t1] = final_df[ch4_t1].astype(float)
        sorted_vals = final_df[ch4_t1].sort_values()
        outliers = list(sorted_vals.head(5)) + list(sorted_vals.tail(5))
        outlier_indices = final_df.sample(n=10).index
        final_df.loc[outlier_indices, ch4_t1] = outliers
        mean_val = final_df[ch4_t1].mean()
        final_df[ch4_t1] = final_df[ch4_t1].apply(lambda x: mean_val if x in outliers else x)

        # Summary statistics
        from scipy.stats import trim_mean, mode
        doc.add_paragraph(f"Mean: {final_df[ch4_t1].mean():.2f}")
        mode_result = mode(final_df[ch4_t1], nan_policy='omit')
        mode_array = mode_result.mode
        mode_value = mode_array[0] if mode_array.size > 0 else "N/A"
        doc.add_paragraph(f"Mode: {mode_value}")
        doc.add_paragraph(f"Median: {final_df[ch4_t1].median():.2f}")
        doc.add_paragraph(f"90% Trimmed Mean: {trim_mean(final_df[ch4_t1], 0.05):.2f}")
        doc.add_paragraph(f"Range: {final_df[ch4_t1].max() - final_df[ch4_t1].min():.2f}")
        doc.add_paragraph(f"Variance: {final_df[ch4_t1].var():.2f}")
        doc.add_paragraph(f"Standard Deviation: {final_df[ch4_t1].std():.2f}")

    # Task 2: Histogram interpretation
    doc.add_heading('Task 2: Histogram Shape', level=2)
    if not show_answer:
        doc.add_paragraph(f"Histogram of variable: {ch4_t1}")
        doc.add_paragraph("Please identify the shape: left-skewed, right-skewed, or normal. Write your interpretation below.")
    else:
        fig, ax = plt.subplots()
        sns.histplot(data=final_df, x=ch4_t1, ax=ax, kde=True)
        ax.set_title(f"Histogram of {ch4_t1}")
        plt.tight_layout()
        hist_path = f"static/plots/histogram_{selected_dataset}.png"
        os.makedirs(os.path.dirname(hist_path), exist_ok=True)
        fig.savefig(hist_path)
        plt.close(fig)
        doc.add_picture(hist_path, width=Inches(5.5))

    # Task 3: Scatterplot
    doc.add_heading('Task 3: Scatterplot of Two Numeric Variables', level=2)
    if not show_answer:
        doc.add_paragraph(f"X: {ch4_scatter_x} | Y: {ch4_scatter_y}")
        x_vals = df[ch4_scatter_x].dropna().astype(str).head(10)
        y_vals = df[ch4_scatter_y].dropna().astype(str).head(10)
        doc.add_paragraph("X values:")
        for val in x_vals:
            doc.add_paragraph(str(val))
        doc.add_paragraph("Y values:")
        for val in y_vals:
            doc.add_paragraph(str(val))
        doc.add_paragraph("Please create a scatterplot and interpret the relationship.")
    else:
        fig, ax = plt.subplots()
        sns.scatterplot(data=final_df, x=ch4_scatter_x, y=ch4_scatter_y, ax=ax)
        ax.set_title("Scatterplot")
        scatter_path = f"static/plots/scatter_{selected_dataset}.png"
        fig.savefig(scatter_path)
        plt.close(fig)
        doc.add_picture(scatter_path, width=Inches(5.5))

    # Task 4: Column chart of two text variables
    doc.add_heading('Task 4: Column Chart from Two Text Variables', level=2)
    if not show_answer:
        doc.add_paragraph(f"X: {ch4_column_x} | Hue: {ch4_column_y}")
        x_vals = df[ch4_column_x].dropna().astype(str).head(10)
        y_vals = df[ch4_column_y].dropna().astype(str).head(10)
        doc.add_paragraph("X values:")
        for val in x_vals:
            doc.add_paragraph(str(val))
        doc.add_paragraph("Y values:")
        for val in y_vals:
            doc.add_paragraph(str(val))
        doc.add_paragraph("Please create a column chart and describe any patterns.")
    else:
        fig, ax = plt.subplots(figsize=(8, 4))
        sns.countplot(data=final_df, x=ch4_column_x, hue=ch4_column_y, ax=ax)
        ax.set_title("Column Chart")
        plt.xticks(rotation=45)
        plt.tight_layout()
        bar_path = f"static/plots/column_chart_{selected_dataset}.png"
        fig.savefig(bar_path)
        plt.close(fig)
        doc.add_picture(bar_path, width=Inches(5.5))


def create_chapter_5_document(doc, selected_dataset, chapter_5_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

def create_chapter_6_document(doc, selected_dataset, chapter_6_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

def create_chapter_7_document(doc, selected_dataset, chapter_7_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

def create_chapter_8_document(doc, selected_dataset, chapter_8_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

def create_chapter_9_document(doc, selected_dataset, chapter_9_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

def create_chapter_10_document(doc, selected_dataset, chapter_10_details, show_answer=False):
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    doc.add_paragraph(f'Selected Dataset: {selected_dataset}')

DOCUMENT_ACTIONS = {
    1: {
        'session_keys': ['selected_dataset', 'chapter_1_details'],
        'action': create_chapter_1_document
    },
    2: {
        'session_keys': ['selected_dataset', 'chapter_2_details'],
        'action': create_chapter_2_document
    },
    3: {
        'session_keys': ['selected_dataset', 'chapter_3_details'],
        'action': create_chapter_3_document
    },
    4: {
        'session_keys': ['selected_dataset', 'chapter_4_details'],
        'action': create_chapter_4_document
    },
    5: {
        'session_keys': ['selected_dataset', 'chapter_5_details'],
        'action': create_chapter_5_document
    },
    6: {
        'session_keys': ['selected_dataset', 'chapter_6_details'],
        'action': create_chapter_6_document
    },
    7: {
        'session_keys': ['selected_dataset', 'chapter_7_details'],
        'action': create_chapter_7_document
    },
    8: {
        'session_keys': ['selected_dataset', 'chapter_8_details'],
        'action': create_chapter_8_document
    },
    9: {
        'session_keys': ['selected_dataset', 'chapter_9_details'],
        'action': create_chapter_9_document
    },
    10: {
        'session_keys': ['selected_dataset', 'chapter_10_details'],
        'action': create_chapter_10_document
    },
}

@chapters.route("/")
def chapters_page():
    session['selected_dataset'] = None
    return render_template('all_chapters.html', chapters_list=ALL_CHAPTERS)


def get_context(chapter, actions_complete=False):
    return {
        'meta': {
            'datasets': Dataset.query.all(),
            'selected_dataset': session.get('selected_dataset', None),
            'chapter_number': chapter,
            'actions_complete': actions_complete
        }
    }


# Chapter 1 Page
@chapters.route('/1', methods=['GET', 'POST'])
def chapter_1():
    chapter_number = 1
    actions_complete = False
    if request.method == 'POST':
        session['chapter_1_details'] = request.form
        actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_1.html', **context)

# Chapter 2 Page
@chapters.route('/2', methods=['GET', 'POST'])
def chapter_2():
    chapter_number = 2
    actions_complete = False
    if request.method == 'POST':
        session['chapter_2_details'] = request.form
        actions_complete = True
    context = get_context(chapter_number, actions_complete)

    db = context.get('meta', {}).get('selected_dataset')
    if db:
        dataset = Dataset.query.filter_by(filename=db).first()
        df = pd.read_json(StringIO(dataset.data)) 
        string_columns = []
        numeric_columns = []

        for col in df.columns:
            # Drop nulls, check if all remaining values are strings
            non_null_values = df[col].dropna()
            if not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, str)).all():
                string_columns.append(col)
            elif not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, (int, float))).all():
                numeric_columns.append(col) 
        context["text_cols"] = string_columns
        context["numeric_cols"] = numeric_columns

    

    return render_template('chapter_2.html', **context)

# Chapter 3 Page
@chapters.route('/3', methods=['GET', 'POST'])
def chapter_3():
    chapter_number = 3
    actions_complete = False

    if request.method == 'POST':
        session['chapter_3_details'] = {**request.form}
        session['chapter_3_details']['plot_types'] = request.form.getlist('plot_types')
        # session['chapter_3_selected_plots'] = selected_plots
        actions_complete = True

    context = get_context(chapter_number, actions_complete)

    # Grab dataset from session/meta and get all column names
    db = context.get('meta', {}).get('selected_dataset')
    if db:
        dataset = Dataset.query.filter_by(filename=db).first()
        df = pd.read_json(StringIO(dataset.data)) 
        string_columns = []
        numeric_columns = []

        for col in df.columns:
            # Drop nulls, check if all remaining values are strings
            non_null_values = df[col].dropna()
            if not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, str)).all():
                string_columns.append(col)
            elif not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, (int, float))).all():
                numeric_columns.append(col) 
        # Just get all columns
        all_columns = df.columns.tolist()
        context["all_columns"] = all_columns
        context["text_cols"] = string_columns
        context["numeric_cols"] = numeric_columns
    context["selected_plots"] = session.get('chapter_3_selected_plots', [])

    return render_template('chapter_3.html', **context)

# Chapter 4 Page
@chapters.route('/4', methods=['GET', 'POST'])
def chapter_4():
    chapter_number = 4
    actions_complete = False
    if request.method == 'POST':
        session['chapter_4_details'] = request.form
        actions_complete = True
    context = get_context(chapter_number, actions_complete)
    db = context.get('meta', {}).get('selected_dataset')
    if db:
        dataset = Dataset.query.filter_by(filename=db).first()
        df = pd.read_json(StringIO(dataset.data)) 
        string_columns = []
        numeric_columns = []

        for col in df.columns:
            # Drop nulls, check if all remaining values are strings
            non_null_values = df[col].dropna()
            if not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, str)).all():
                string_columns.append(col)
            elif not non_null_values.empty and non_null_values.apply(lambda x: isinstance(x, (int, float))).all():
                numeric_columns.append(col) 
        # Just get all columns
        all_columns = df.columns.tolist()
        context["all_columns"] = all_columns
        context["text_cols"] = string_columns
        context["numeric_cols"] = numeric_columns
    return render_template('chapter_4.html', **context)

# Chapter 5 Page
@chapters.route('/5')
def chapter_5():
    chapter_number = 5
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_5.html', **context)

# Chapter 6 Page
@chapters.route('/6')
def chapter_6():
    chapter_number = 6
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_6.html', **context)

# Chapter 7 Page
@chapters.route('/7')
def chapter_7():
    chapter_number = 7
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_7.html', **context)

# Chapter 8 Page
@chapters.route('/8')
def chapter_8():
    chapter_number = 8
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_8.html', **context)

# Chapter 9 Page
@chapters.route('/9')
def chapter_9():
    chapter_number = 9
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_9.html', **context)

# Chapter 10 Page
@chapters.route('/10')
def chapter_10():
    chapter_number = 10
    actions_complete = True
    context = get_context(chapter_number, actions_complete)
    return render_template('chapter_10.html', **context)

# Confirm Dataset for Chapter 2
@chapters.route('/<int:index>/dataset/confirm', methods=['POST'])
def confirm_dataset_for_chapter(index):
    session['selected_dataset'] = request.form.get('selected_dataset')
    return redirect(f'/chapters/{index}')

# Generate Word Document for Chapter 2 (Restored to Plain Text Paths)
@chapters.route('/<int:index>/generate/word/<string:q_key>')
def generate_word_for_chapter(index, q_key):
    selected_dataset = session.get('selected_dataset')
    procedure = DOCUMENT_ACTIONS.get(index)
    if not selected_dataset:
        flash("⚠ Please select a dataset first.", "warning")
        return redirect(f'/chapters/{index}')
    if not procedure:
        return redirect(f'/chapters/{index}')

    filename = f"Chapter_{index}_Assignment_{q_key.capitalize()}.docx"

    # Create Word Document
    word_path = os.path.join("static", "generated", filename)

    doc = Document()
    # doc.add_heading(f'Chapter {index} Assignment', level=1)

    props = {
        "show_answer": q_key == 'answer'
    }
    for key in procedure.get('session_keys', []):
        props[key] = session.get(key)
    print(procedure, props)
    procedure['action'](doc, **props)

    doc.save(word_path)

    return send_file(word_path, as_attachment=True, download_name=filename)
