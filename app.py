import os
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, send_file, session
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from docx import Document

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///datasets.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = "supersecretkey"

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)

# Define Dataset Model
class Dataset(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(100), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    data = db.Column(db.Text, nullable=False)

# Initialize Database
with app.app_context():
    db.create_all()

# Home Page
@app.route('/')
def index():
    dataset_exists = Dataset.query.first() is not None
    return render_template('index.html', dataset_exists=dataset_exists)

@app.route('/chapters')
def chapters():
    return render_template('chapters.html')


# Upload Dataset
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'dataset' not in request.files:
        flash("❌ No file selected!", "error")
        return redirect(url_for('index'))

    file = request.files['dataset']
    if file.filename == '':
        flash("❌ No file selected!", "error")
        return redirect(url_for('index'))

    if file and file.filename.endswith('.csv'):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)

        # Prevent duplicate uploads
        existing_dataset = Dataset.query.filter_by(filename=filename).first()
        if existing_dataset:
            flash("⚠ Dataset already exists. Please use the existing one.", "warning")
            return redirect(url_for('index'))

        file.save(filepath)

        df = pd.read_csv(filepath)
        dataset_entry = Dataset(title=filename, filename=filename, data=df.to_json())
        db.session.add(dataset_entry)
        db.session.commit()

        flash("✅ Dataset uploaded successfully!", "success")
        return redirect(url_for('index'))
    else:
        flash("❌ Invalid file format. Please upload a CSV file.", "error")
        return redirect(url_for('index'))

# Confirm Dataset for Chapter 2
@app.route('/confirm_dataset_chapter2', methods=['POST'])
def confirm_dataset_chapter2():
    selected_dataset = request.form.get('selected_dataset')
    session['selected_dataset'] = selected_dataset
    return redirect(url_for('chapter_2'))

# Chapter 2 Page
@app.route('/chapter/2')
def chapter_2():
    datasets = Dataset.query.all()
    selected_dataset = session.get('selected_dataset', None)
    return render_template('chapter_2.html', datasets=datasets, selected_dataset=selected_dataset)

# Generate Word Document for Chapter 2 (Restored to Plain Text Paths)
@app.route('/generate_word/chapter/2')
def generate_word_chapter_2():
    selected_dataset = session.get('selected_dataset')
    if not selected_dataset:
        flash("⚠ Please select a dataset first.", "warning")
        return redirect(url_for('chapter_2'))

    # Load dataset
    dataset = Dataset.query.filter_by(filename=selected_dataset).first()
    df = pd.read_json(dataset.data)

    # Duplicate 25 random rows
    duplicated_df = pd.concat([df, df.sample(25, replace=True)])
    duplicated_path = f"static/csv/duplicated_{selected_dataset}"
    original_path = f"static/csv/original_{selected_dataset}"

    duplicated_df.to_csv(duplicated_path, index=False)
    df.to_csv(original_path, index=False)

    # Create Word Document
    word_path = "static/generated/Chapter_2_Assignment.docx"
    doc = Document()
    doc.add_heading('Chapter 2 Assignment', level=1)

    doc.add_heading('Task 1: Remove Duplicate Records', level=2)
    doc.add_paragraph("Download the dataset with duplicated rows:")

    # ✅ Plain text links for downloading CSV files
    doc.add_paragraph(f"📂 Duplicated Dataset: {os.path.abspath(duplicated_path)}")
    doc.add_paragraph("\n✅ Answer Key:")
    doc.add_paragraph(f"📂 Original Dataset: {os.path.abspath(original_path)}")

    # Save the Word document
    doc.save(word_path)

    return send_file(word_path, as_attachment=True, download_name="Chapter_2_Assignment.docx")

if __name__ == '__main__':
    app.run(debug=True)
